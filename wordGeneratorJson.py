from jsonConstructor import InformationProcessorJson
from docx import Document
import os
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
class WordGeneratorJson:
    def __init__(self,WordPath,JsonInoformation:InformationProcessorJson,information) -> None:
        self._WordPath=WordPath
        self._Document=Document()
        self._JsonInoformation=JsonInoformation
        self._JsonInoformation.fillByJsonObject(information)
        self.generateWord()
        pass
    def generateWord(self):
        self._Document.add_heading(str(self._JsonInoformation.Company.Name), level=1)

        # Add company details as bullet points
        self._Document.add_paragraph("Adresse : "+str(self._JsonInoformation.Company.Address))
        self._Document.add_paragraph("MF : "+str(self._JsonInoformation.Company.MF))
        self._Document.add_paragraph("Tél : "+str(self._JsonInoformation.Company.Phone))

        # Add the client information
        self._Document.add_paragraph('Coordonnées Client, Nom : '+str(self._JsonInoformation.Client.Name), style='List Bullet')
        self._Document.add_paragraph('Adresse : '+str(self._JsonInoformation.Client.Address), style='List Bullet')
        self._Document.add_paragraph('Tél : '+str(self._JsonInoformation.Client.Phone), style='List Bullet')
        currentDate=datetime.now()
        self._Document.add_paragraph('Date : '+str(currentDate.year)+'\\'+str(currentDate.month)+'\\'+str(currentDate.day)).alignment = WD_ALIGN_PARAGRAPH.RIGHT

        RowNumber=len(self._JsonInoformation.Products)+1
        self._Document.add_heading('Devis N°'+str(self._JsonInoformation.ReceiptNumber), level=2)
        table = self._Document.add_table(rows=RowNumber, cols=5)
        table.style = 'Table Grid'

        # Fill in the table header
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Désignation'
        header_cells[1].text = 'Unité'
        header_cells[2].text = 'Quantité'
        header_cells[3].text = 'Prix unitaire HT'
        header_cells[4].text = 'Total'
        for i in range(1,RowNumber):
            row_cells = table.rows[i].cells
            row_cells[0].text = str(self._JsonInoformation.Products[i-1].Name)
            row_cells[1].text = 'P'
            if self._JsonInoformation.Products[i-1].Quantity==1:
                row_cells[2].text ='-'
            else:
                row_cells[2].text = str(self._JsonInoformation.Products[i-1].Quantity)
            row_cells[3].text = str(self._JsonInoformation.Products[i-1].PricePerUnit)
            row_cells[4].text = str(self._JsonInoformation.Products[i-1].Total)

        row = table.add_row()
        row_cells = row.cells
        row_cells[3].text = 'Total HT'
        row_cells[4].text = str(self._JsonInoformation.Total())

        self._Document.add_paragraph('Modalité de paiement : 100%')
        self._Document.save(os.path.join(self._WordPath,'Devis_'+str(self._JsonInoformation.ReceiptNumber)+'.docx'))