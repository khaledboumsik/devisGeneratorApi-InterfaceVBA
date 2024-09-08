from inputConstructor import InputConstructor
from docx import Document
import os
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
class WordGeneratorPrompt:
    def __init__(self,WordPath,InputConstuctor:InputConstructor) -> None:
        self._WordPath=WordPath
        self._Document=Document()
        self._InputConstuctor=InputConstuctor
        self.generateWord()
        pass
    def generateWord(self):
        ReceiptNumber=int(input("What is the number of this receipt"))
        self._Document.add_heading(str(self._InputConstuctor._companyHandler._Company.Name), level=1)

        # Add company details as bullet points
        self._Document.add_paragraph("Adresse : "+str(self._InputConstuctor._companyHandler._Company.Address))
        self._Document.add_paragraph("MF : "+str(self._InputConstuctor._companyHandler._Company.MF))
        self._Document.add_paragraph("Tél : "+str(self._InputConstuctor._companyHandler._Company.Phone))

        # Add the client information
        self._Document.add_paragraph('Coordonnées Client, Nom : '+str(self._InputConstuctor._clientHandler._Client.Name), style='List Bullet')
        self._Document.add_paragraph('Adresse : '+str(self._InputConstuctor._clientHandler._Client.Address), style='List Bullet')
        self._Document.add_paragraph('Tél : '+str(self._InputConstuctor._clientHandler._Client.Phone), style='List Bullet')
        currentDate=datetime.now()
        self._Document.add_paragraph('Date : '+str(currentDate.year)+'\\'+str(currentDate.month)+'\\'+str(currentDate.day)).alignment = WD_ALIGN_PARAGRAPH.RIGHT

        RowNumber=len(self._InputConstuctor._ProductHandler._Products)+1
        self._Document.add_heading('Devis N°'+str(ReceiptNumber), level=2)
        table = self._Document.add_table(rows=RowNumber, cols=5)
        table.style = 'Table Grid'

        # Fill in the table header
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Désignation'
        header_cells[1].text = 'Unité'
        header_cells[2].text = 'Quantité'
        header_cells[3].text = 'Prix unitaire HT'
        header_cells[4].text = 'Total'
        for i in range(1,RowNumber):
            row_cells = table.rows[i].cells
            row_cells[0].text = str(self._InputConstuctor._ProductHandler._Products[i-1].Name)
            row_cells[1].text = 'P'
            if self._InputConstuctor._ProductHandler._Products[i-1].Quantity==1:
                row_cells[2].text ='-'
            else:
                row_cells[2].text = str(self._InputConstuctor._ProductHandler._Products[i-1].Quantity)
            row_cells[3].text = str(self._InputConstuctor._ProductHandler._Products[i-1].PricePerUnit)
            row_cells[4].text = str(self._InputConstuctor._ProductHandler._Products[i-1].Total)

        row = table.add_row()
        row_cells = row.cells
        row_cells[3].text = 'Total HT'
        self._InputConstuctor._ProductHandler.calculateGlobalTotal()
        row_cells[4].text = str(self._InputConstuctor._ProductHandler.Total)

        self._Document.add_paragraph('Modalité de paiement : 100%')
        self._Document.save(os.path.join(self._WordPath,'Devis_'+str(ReceiptNumber)+'.docx'))